import argparse
import json
import re
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")


def slugify(value, suffix):
    value = re.sub(r"[\\/:*?\"<>|]+", "-", value or "study-artifact")
    value = re.sub(r"\s+", " ", value).strip().strip(".")
    value = value.strip("- ")
    return f"{value or 'study-artifact'}{suffix}"


def default_output_path(title, suffix):
    output_dir = Path.cwd() / "courseware-output"
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / slugify(title, suffix)


def iter_markdown_blocks(markdown):
    in_code = False
    code_lines = []
    code_lang = ""
    for raw_line in (markdown or "").splitlines():
        line = raw_line.rstrip()
        if line.strip().startswith("```"):
            if in_code:
                yield {"type": "code", "language": code_lang, "text": "\n".join(code_lines)}
                in_code = False
                code_lines = []
                code_lang = ""
            else:
                in_code = True
                code_lang = line.strip().strip("`").strip()
            continue
        if in_code:
            code_lines.append(line)
            continue
        yield {"type": "line", "text": line}
    if in_code and code_lines:
        yield {"type": "code", "language": code_lang, "text": "\n".join(code_lines)}


def strip_outer_markdown_fence(markdown):
    text = (markdown or "").strip()
    match = re.match(r"^```(?:markdown|md)?\s*(.*?)```$", text, flags=re.DOTALL | re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return markdown or ""


def normalize_markdown(markdown):
    text = strip_outer_markdown_fence(markdown)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = []
    for raw in text.split("\n"):
        line = raw.rstrip()
        stripped = line.strip()
        if re.match(r"^[一二三四五六七八九十]+[、.]\s*", stripped):
            lines.append("## " + re.sub(r"^[一二三四五六七八九十]+[、.]\s*", "", stripped))
            continue
        if re.match(r"^第[一二三四五六七八九十0-9]+[章节部分]\s*[:：]?\s*", stripped):
            lines.append("## " + stripped)
            continue
        if re.match(r"^(核心概念|重点|难点|总结|复习建议|学习建议|知识结构|主线)[:：]\s*$", stripped):
            lines.append("## " + stripped.rstrip(":："))
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def set_style_font(style, name="宋体", size_pt=12, bold=None, color=None):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    style.font.name = name
    style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)


def set_run_font(run, name="宋体", size_pt=12, bold=None):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)


def configure_doc(doc):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    set_style_font(styles["Normal"], "宋体", 12)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 12, RGBColor(31, 78, 121)),
        ("Heading 2", 12, RGBColor(47, 84, 150)),
        ("Heading 3", 12, RGBColor(68, 68, 68)),
    ]:
        style = styles[name]
        set_style_font(style, "黑体", size, bold=True, color=color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    title_style = styles.add_style("Study Title", 1)
    set_style_font(title_style, "黑体", 12, bold=True, color=RGBColor(31, 78, 121))
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(6)

    subtitle_style = styles.add_style("Study Subtitle", 1)
    set_style_font(subtitle_style, "宋体", 12, color=RGBColor(89, 89, 89))
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(14)

    code_style = styles.add_style("Code Block", 1)
    set_style_font(code_style, "宋体", 12)
    code_style.paragraph_format.left_indent = Inches(0.25)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(6)

    for list_style_name in ("List Bullet", "List Number"):
        if list_style_name in styles:
            set_style_font(styles[list_style_name], "宋体", 12)


def add_inline_markdown(paragraph, text):
    from docx.shared import Pt

    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, "宋体", 12, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "宋体", 12)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, "宋体", 12)


def is_table_separator(line):
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def is_table_row(line):
    return "|" in line and not line.strip().startswith("```")


def is_box_table_line(line):
    return bool(re.search(r"[┌┬┐├┼┤└┴┘│─]", line))


def split_table_row(line):
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return cells


def add_box_table(doc, table_lines):
    for line in table_lines:
        p = doc.add_paragraph(style="Code Block")
        run = p.add_run(line)
        set_run_font(run, "宋体", 12)


def add_markdown_table(doc, table_lines):
    rows = [split_table_row(line) for line in table_lines if not is_table_separator(line)]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for col_index in range(column_count):
            value = row_values[col_index] if col_index < len(row_values) else ""
            paragraph = cells[col_index].paragraphs[0]
            add_inline_markdown(paragraph, value)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True


def add_markdown_to_doc(doc, markdown):
    from docx.shared import Inches, Pt

    pending_number = 1
    blocks = list(iter_markdown_blocks(normalize_markdown(markdown)))
    index = 0
    while index < len(blocks):
        block = blocks[index]
        if block["type"] == "code":
            for line in block["text"].splitlines() or [""]:
                p = doc.add_paragraph(style="Code Block")
                run = p.add_run(line)
                set_run_font(run, "宋体", 12)
            pending_number = 1
            index += 1
            continue

        line = block["text"]
        stripped = line.strip()
        if not stripped:
            pending_number = 1
            index += 1
            continue

        if is_box_table_line(stripped):
            table_lines = []
            while index < len(blocks) and blocks[index]["type"] == "line" and is_box_table_line(blocks[index]["text"].strip()):
                table_lines.append(blocks[index]["text"].rstrip())
                index += 1
            add_box_table(doc, table_lines)
            pending_number = 1
            continue

        if is_table_row(stripped):
            table_lines = []
            while index < len(blocks) and blocks[index]["type"] == "line" and is_table_row(blocks[index]["text"].strip()):
                table_lines.append(blocks[index]["text"].strip())
                index += 1
            add_markdown_table(doc, table_lines)
            pending_number = 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif re.match(r"^[-*]\s+", stripped):
            indent_level = max(0, (len(line) - len(line.lstrip(" "))) // 2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + 0.18 * indent_level)
            add_inline_markdown(p, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+[.)]\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, re.sub(r"^\d+[.)]\s+", "", stripped))
            pending_number += 1
        else:
            p = doc.add_paragraph()
            add_inline_markdown(p, stripped)
        index += 1


def create_study_docx(payload):
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError(f"python-docx is required to create DOCX files: {exc}") from exc

    title = payload.get("title") or "Courseware Study Notes"
    subtitle = payload.get("subtitle") or "Structured summary, review points, and exam preparation notes"
    content = payload.get("content_markdown") or ""
    mindmap = payload.get("mindmap_mermaid") or ""
    output_path = Path(payload.get("output_path") or default_output_path(title, ".docx"))
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_doc(doc)
    doc.add_paragraph(title, style="Study Title")
    if subtitle:
        doc.add_paragraph(subtitle, style="Study Subtitle")
    add_markdown_to_doc(doc, content)

    if mindmap.strip():
        doc.add_page_break()
        doc.add_heading("Mermaid Mind Map Source", level=1)
        add_markdown_to_doc(doc, f"```mermaid\n{extract_mermaid(mindmap)}\n```")

    doc.core_properties.title = title
    doc.core_properties.subject = "Courseware study notes"
    doc.save(output_path)
    return {
        "type": "docx",
        "path": str(output_path.resolve()),
        "title": title,
        "bytes": output_path.stat().st_size,
    }


def extract_mermaid(value):
    text = (value or "").strip()
    match = re.search(r"```(?:mermaid)?\s*(.*?)```", text, flags=re.DOTALL | re.IGNORECASE)
    if match:
        text = match.group(1).strip()
    return text


def create_mind_map_file(payload):
    title = payload.get("title") or "Courseware Mind Map"
    mermaid = extract_mermaid(payload.get("mermaid_mindmap") or "")
    if not mermaid:
        raise ValueError("mermaid_mindmap is required.")
    if not mermaid.lstrip().startswith("mindmap"):
        mermaid = "mindmap\n" + mermaid

    file_format = (payload.get("format") or "md").lower()
    suffix = ".mmd" if file_format == "mmd" else ".md"
    output_path = Path(payload.get("output_path") or default_output_path(title, suffix))
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if suffix == ".mmd":
        content = mermaid.rstrip() + "\n"
    else:
        content = f"# {title}\n\n```mermaid\n{mermaid.rstrip()}\n```\n"
    output_path.write_text(content, encoding="utf-8")
    return {
        "type": "mindmap",
        "path": str(output_path.resolve()),
        "title": title,
        "format": file_format,
        "bytes": output_path.stat().st_size,
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--mode", required=True, choices=["docx", "mindmap"])
    args = parser.parse_args()
    payload = json.load(sys.stdin)
    if args.mode == "docx":
        result = create_study_docx(payload)
    else:
        result = create_mind_map_file(payload)
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        sys.exit(1)
